from docx import Document
import os

OUTPUT_DIR = os.path.join(os.path.dirname(__file__))
OUTPUT_PATH = os.path.join(OUTPUT_DIR, 'data_dictionary.docx')

TABLES = [
    ('App_Users', [
        ('id', 'INT UNSIGNED AUTO_INCREMENT', 'PK; AUTO_INCREMENT', 'Internal user identifier.'),
        ('username', 'VARCHAR(80)', 'NOT NULL; UNIQUE', 'User login name.'),
        ('password_hash', 'VARCHAR(255)', 'NOT NULL', 'Hashed password (bcrypt).'),
        ('created_at', 'TIMESTAMP', 'NOT NULL; DEFAULT CURRENT_TIMESTAMP', 'Account creation timestamp.'),
    ]),
    ('Tickers', [
        ('id', 'INT UNSIGNED AUTO_INCREMENT', 'PK; AUTO_INCREMENT', 'Internal ticker id.'),
        ('symbol', 'VARCHAR(20)', 'NOT NULL; UNIQUE', 'Stock symbol (e.g., AAPL).'),
        ('slug', 'VARCHAR(40)', 'NOT NULL; UNIQUE', 'URL-friendly identifier.'),
        ('display_name', 'VARCHAR(120)', 'NOT NULL', 'Human-readable company/ticker name.'),
        ('created_at', 'TIMESTAMP', 'NOT NULL; DEFAULT CURRENT_TIMESTAMP', 'Record creation time.'),
    ]),
    ('Prediction_Jobs', [
        ('id', 'INT UNSIGNED AUTO_INCREMENT', 'PK; AUTO_INCREMENT', 'Job id.'),
        ('user_id', 'INT UNSIGNED', "NULL; FK → app_users(id) ON DELETE SET NULL", 'Owner (nullable for anonymous/demo jobs).'),
        ('ticker_id', 'INT UNSIGNED', 'NOT NULL; FK → tickers(id) ON DELETE CASCADE', 'Related ticker.'),
        ('requested_ticker', 'VARCHAR(20)', 'NOT NULL', 'Requested ticker symbol (copied into job).'),
        ('status', "ENUM('queued','running','completed','failed')", "NOT NULL; DEFAULT 'queued'", 'Job lifecycle state.'),
        ('requested_device', 'VARCHAR(20)', "NOT NULL; DEFAULT 'gpu'", 'Device preference (gpu/cpu).'),
        ('requested_epochs', 'INT UNSIGNED', 'NOT NULL; DEFAULT 40', 'Training epochs requested.'),
        ('requested_batch_size', 'INT UNSIGNED', 'NOT NULL; DEFAULT 16', 'Batch size requested.'),
        ('requested_prediction_days', 'INT UNSIGNED', 'NOT NULL; DEFAULT 30', 'Days to predict.'),
        ('requested_future_days', 'INT UNSIGNED', 'NOT NULL; DEFAULT 30', 'Horizon for future data used.'),
        ('requested_mc_runs', 'INT UNSIGNED', 'NOT NULL; DEFAULT 100', 'Monte Carlo runs for uncertainty estimation.'),
        ('output_message', 'TEXT', 'NULL', 'Output or log text from run.'),
        ('failure_message', 'TEXT', 'NULL', 'Error details if failed.'),
        ('remote_manifest_json', 'LONGTEXT', 'NULL', 'Remote job manifest (JSON).'),
        ('created_at', 'TIMESTAMP', 'NOT NULL; DEFAULT CURRENT_TIMESTAMP', 'Job creation time.'),
        ('started_at', 'TIMESTAMP', 'NULL', 'When job started.'),
        ('completed_at', 'TIMESTAMP', 'NULL', 'When job completed.'),
        ('Indexes/Notes', '—', '—', "INDEX (status, created_at); INDEX (ticker_id, created_at)"),
    ]),
    ('Saved_Graphs', [
        ('id', 'INT UNSIGNED AUTO_INCREMENT', 'PK; AUTO_INCREMENT', 'Saved graph id.'),
        ('job_id', 'INT UNSIGNED', "NOT NULL; UNIQUE; FK → prediction_jobs(id) ON DELETE CASCADE", 'Linked job (unique → 1:1 relationship).'),
        ('user_id', 'INT UNSIGNED', 'NULL; FK → app_users(id) ON DELETE SET NULL', 'Owning user (nullable).'),
        ('ticker_id', 'INT UNSIGNED', 'NOT NULL; FK → tickers(id) ON DELETE CASCADE', 'Ticker shown.'),
        ('title', 'VARCHAR(180)', 'NOT NULL', 'Chart title.'),
        ('summary_text', 'TEXT', 'NULL', 'Short summary / caption.'),
        ('summary_plot_path', 'VARCHAR(255)', 'NULL', 'File path for summary image.'),
        ('detail_plot_path', 'VARCHAR(255)', 'NULL', 'File path for detailed image.'),
        ('predictions_csv_path', 'VARCHAR(255)', 'NULL', 'CSV of predicted values.'),
        ('forecast_csv_path', 'VARCHAR(255)', 'NULL', 'CSV of forecast / inputs.'),
        ('remote_job_directory', 'VARCHAR(255)', 'NULL', 'Remote storage path for job artifacts.'),
        ('created_at', 'TIMESTAMP', 'NOT NULL; DEFAULT CURRENT_TIMESTAMP', 'When saved.'),
        ('Indexes/Notes', '—', '—', 'INDEX (ticker_id, created_at) for queries.'),
    ]),
    ('Saved_Graph_Assets', [
        ('id', 'INT UNSIGNED AUTO_INCREMENT', 'PK; AUTO_INCREMENT', 'Asset id.'),
        ('graph_id', 'INT UNSIGNED', 'NOT NULL; FK → saved_graphs(id) ON DELETE CASCADE', 'Parent graph.'),
        ('asset_kind', "ENUM('summary','detail','residuals')", 'NOT NULL', 'Type of asset.'),
        ('mime_type', 'VARCHAR(120)', 'NOT NULL', 'MIME type (e.g., image/png).'),
        ('original_name', 'VARCHAR(255)', 'NOT NULL', 'Original filename at upload.'),
        ('binary_data', 'LONGBLOB', 'NOT NULL', 'Binary file contents.'),
        ('created_at', 'TIMESTAMP', 'NOT NULL; DEFAULT CURRENT_TIMESTAMP', 'Upload time.'),
        ('Constraints/Notes', '—', '—', 'UNIQUE (graph_id, asset_kind) — one asset per kind per graph.'),
    ]),
]


def make_doc(output_path):
    doc = Document()
    doc.add_heading('Data Dictionary', level=1)
    for table_name, rows in TABLES:
        doc.add_heading(table_name, level=2)
        table = doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Fieldname'
        hdr_cells[1].text = 'Data type'
        hdr_cells[2].text = 'Validation'
        hdr_cells[3].text = 'Description'
        for field, dtype, validation, desc in rows:
            r = table.add_row().cells
            r[0].text = str(field)
            r[1].text = str(dtype)
            r[2].text = str(validation)
            r[3].text = str(desc)
        doc.add_paragraph('')
    doc.save(output_path)


if __name__ == '__main__':
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    make_doc(OUTPUT_PATH)
    print('Wrote:', OUTPUT_PATH)
